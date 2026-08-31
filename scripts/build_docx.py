"""Render the solution walkthrough to .docx.

Submission requirement #2 is "a word document (as .docx)". Keeping the source in Markdown
and generating the .docx means the text stays reviewable and diffable in git, and the
deliverable is reproducible rather than a binary somebody hand-edited once.

    python scripts/build_docx.py

Writes docs/submission/solution-walkthrough.docx.

Deliberately a small subset of Markdown -- headings, paragraphs, bullets, fenced code,
pipe tables, blockquotes and bold/inline-code spans. That is what the document uses. It is
not a general converter and should not become one.
"""

from __future__ import annotations

import re
import sys
import os
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from adversarial_payments.config import ROOT

SOURCE = ROOT / "docs" / "submission" / "solution-walkthrough.md"
# Overridable so a rebuild is possible while the document is open in Word, which holds an
# exclusive lock. Building beside it and swapping afterwards beats asking someone to close a
# file they are in the middle of reading.
DEST = Path(
    os.environ.get("DOCX_OUT", ROOT / "docs" / "submission" / "solution-walkthrough.docx")
)

ACCENT = RGBColor(0x1F, 0x3A, 0x5F)
MUTED = RGBColor(0x60, 0x6A, 0x78)
PENDING = RGBColor(0xB4, 0x54, 0x09)


def _style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size in (("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(16)
        style.paragraph_format.space_after = Pt(6)


# `code`, **bold**, and [[PENDING]] markers all need distinct runs.
_SPAN = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[\[[^\]]+\]\])")


def _add_runs(paragraph, text: str) -> None:
    """Split inline markup into runs. Nothing here nests, so one pass is enough."""
    for piece in _SPAN.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            paragraph.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif piece.startswith("[[") and piece.endswith("]]"):
            # Unresolved figures must be impossible to miss on a printed page.
            run = paragraph.add_run(piece)
            run.bold = True
            run.font.color.rgb = PENDING
        else:
            paragraph.add_run(piece)


def _table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            if c >= len(table.columns):
                continue
            cell = table.cell(r, c)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_runs(paragraph, cell_text)
            if r == 0:
                for run in paragraph.runs:
                    run.bold = True
    doc.add_paragraph()


def _is_separator(cells: list[str]) -> bool:
    return all(set(c.strip()) <= {"-", ":"} and c.strip() for c in cells)


_IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)$")


def convert(md: str) -> Document:
    doc = Document()
    _style(doc)

    lines = md.split("\n")
    i = 0
    pending_table: list[list[str]] = []

    def flush_table() -> None:
        nonlocal pending_table
        if pending_table:
            _table(doc, pending_table)
            pending_table = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            if not _is_separator(cells):
                pending_table.append(cells)
            i += 1
            continue
        flush_table()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # ![caption](figures/x.png) -- its own paragraph, never inline.
        img = _IMAGE_RE.match(stripped)
        if img:
            caption, rel = img.group(1), img.group(2)
            path = (ROOT / "docs" / "submission" / rel).resolve()
            if path.exists():
                doc.add_picture(str(path), width=Inches(6.0))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if caption:
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cap.add_run(caption)
                    run.italic = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x6B, 0x65, 0x58)
            else:
                # A missing figure must be loud, exactly as an unresolved marker is.
                warn = doc.add_paragraph()
                r = warn.add_run(f"[MISSING FIGURE: {rel}]")
                r.bold = True
                r.font.color.rgb = RGBColor(0xC9, 0x37, 0x2B)
                print(f"  WARNING: missing figure {rel}", file=sys.stderr)
            i += 1
            continue

        if stripped.startswith("```"):
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            i += 1
            paragraph = doc.add_paragraph()
            run = paragraph.add_run("\n".join(block))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            paragraph.paragraph_format.left_indent = Pt(18)
            continue

        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            if level == 1 and not doc.paragraphs:
                heading = doc.add_heading(text, 0)
                heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, min(level, 4))
            i += 1
            continue

        if stripped.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            _add_runs(paragraph, stripped[2:])
            for run in paragraph.runs:
                run.italic = True
                if run.font.color.rgb is None:
                    run.font.color.rgb = MUTED
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_runs(paragraph, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s", stripped):
            paragraph = doc.add_paragraph(style="List Number")
            _add_runs(paragraph, re.sub(r"^\d+\.\s", "", stripped))
            i += 1
            continue

        # Consecutive non-blank lines are one wrapped paragraph.
        block = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#|>|\||```|- |\* |\d+\.\s|---)", lines[i].strip()
        ):
            block.append(lines[i].strip())
            i += 1
        paragraph = doc.add_paragraph()
        _add_runs(paragraph, " ".join(block))

    flush_table()
    return doc


def main() -> int:
    if not SOURCE.exists():
        print(f"missing source: {SOURCE}")
        return 1

    md = SOURCE.read_text(encoding="utf-8")
    doc = convert(md)
    DEST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DEST)

    pending = md.count("[[")
    print(f"wrote {DEST.relative_to(ROOT)}  ({DEST.stat().st_size:,} bytes)")
    print(f"  paragraphs: {len(doc.paragraphs)}   tables: {len(doc.tables)}")
    if pending:
        print(f"  {pending} [[PENDING]] marker(s) -- these render in orange and must be")
        print("  resolved or removed before submission.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
